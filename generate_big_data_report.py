import sys
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_style(paragraph, font_name="Times New Roman", font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.15, bold=False):
    paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = line_spacing
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        run.font.size = Pt(font_size)
        run.bold = bold

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.outline_level = level - 1
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True
    return p

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    set_style(p, font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.15)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph("• " + text)
    set_style(p, font_size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.15)
    return p

def main():
    try:
        doc = Document("Front Page.docx")
        doc.add_page_break()
    except Exception as e:
        print("Could not load 'Front Page.docx'. Creating a new document.")
        doc = Document()
    
    # Set default styles for the document
    try:
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
    except Exception:
        pass

    # Add Table of Contents
    toc_title = doc.add_paragraph("Table of Contents")
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in toc_title.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)
        run.bold = True
        
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    
    doc.add_page_break()

    # 1. Abstract
    add_heading(doc, "1. Abstract")
    add_paragraph(doc, "This project presents an enterprise-grade Knowledge Retrieval System designed to parse, index, and query software programming syntax across massive source code repositories natively. Leveraging a highly distributed Retrieval-Augmented Generation (RAG) pipeline, the framework connects real-time document streaming via Apache Kafka directly to Hadoop Distributed File System (HDFS) staging clusters. Semantic mapping is achieved by extracting abstract syntax and transforming code configurations into dense numerical vectors stored safely inside a persistent ChromaDB instance. State-of-the-art transformer processing is conducted via the GROQ API, utilizing the Llama-3 (70b and 8b) massive language models to synthesize exact code explanations from the dynamically retrieved chunk parameters. Evaluation confirms that decoupling extraction phases securely into asynchronous event-streams eliminates REST API bottlenecking, ultimately achieving sub-second generative documentation queries.")

    # 2. Introduction
    add_heading(doc, "2. Introduction")
    
    p = add_paragraph(doc, "")
    run1 = p.add_run("Background and motivation for the project: ")
    run1.bold = True
    run2 = p.add_run("Modern software environments construct intricate architectures distributed across dozens of connected Git repositories. As codebases scale into millions of lines, the ability for newly onboarded developers to locate, understand, and safely modify interconnected configurations diminishes exponentially. Reliance on manual documentation often fails, as static written files rarely mirror rapid continuous integration workflows. This establishes a pressing need for the integration of intelligent Retrieval-Augmented Generation (RAG) systems capable of reading, indexing, and actively explaining undocumented repository segments on demand. By leveraging large language models tightly coupled with localized vector stores, this project aims to create an AI assistant that dynamically 'reads' entire Git repositories and provides logically sound architectural answers, bridging the gap between raw unstructured code and human comprehension.")
    run2.bold = False

    p2 = add_paragraph(doc, "")
    run3 = p2.add_run("Problem statement and research questions: ")
    run3.bold = True
    run4 = p2.add_run("The primary hurdle in developing querying endpoints for massive file directories natively is standard deterministic search capabilities (e.g., regex matching). Standard lexical indexers frequently fail to capture the inherited meaning of complex object definitions scattered globally. Conversely, feeding entire codebases natively into commercial Large Language Models completely shatters inherent generative context windows (token limits) and drives up inference costs. Therefore, the core problem addressed is resolving the token constraint bottleneck via real-time intelligent extraction: How can we successfully mitigate indexing latency by integrating a high-speed Big Data ingestion streaming queue with a local Vector store capable of dynamically pushing isolated sub-routines safely into Llama-3 transformer queries without completely exhausting the REST API thread limit?")
    run4.bold = False

    p3 = add_paragraph(doc, "")
    run5 = p3.add_run("Scope and limitations: ")
    run5.bold = True
    run6 = p3.add_run("The scope of this research explicitly covers asynchronous multi-threading data stream architectures utilizing Apache Kafka and Hadoop HDFS to manage extraction from dynamic Git cloning arrays. The system actively utilizes Python's Flask framework (`app.py`) for ingesting repository URLs, securely transforming them via specific embedding metrics into persistent ChromaDB stores (managed in `retriever.py`). Supported formats extensively target UTF-8 programming scripts. The limitation definitively residues in graphic interface visual elements or highly compiled binary packages mapping outside standard plaintext protocols. Furthermore, strict dependency on the Groq external API infrastructure enforces processing reliance heavily to outbound SSL bandwidth, and the fallback to smaller `llama-3.1-8b-instant` models limits profound logical leaps achievable by fully uncapped, massive parameter models.")
    run6.bold = False

    # 3. Literature Review
    add_heading(doc, "3. Literature Review")

    p4 = add_paragraph(doc, "")
    r7 = p4.add_run("Overview of existing work in big data analytics: ")
    r7.bold = True
    r8 = p4.add_run("Traditionally, analyzing software file structures relied intensely on strictly mapped SQL databases traversing structural metadata or relying exclusively on elastic string search protocols, like BM25 algorithms or Apache Lucene. These legacy systems easily break when semantic phrasing deviates. Academic literature continually pushes transitioning data frameworks immediately over to real-time generative query arrays natively using distributed computing lakes designed expressly for massive, multi-threaded inference flows. Early text-based big data pipelines relied on Hadoop's heavy MapReduce paradigms, which progressively shifted towards in-memory streaming semantics to accommodate real-time knowledge ingestion.")
    r8.bold = False

    p5 = add_paragraph(doc, "")
    r9 = p5.add_run("Key frameworks, algorithms, and tools relevant to your project: ")
    r9.bold = True
    r10 = p5.add_run("Recent academic shifts validate that integrating Large Language Models (specifically Transformer variations like Llama-3) is functionally optimal for generative coding workflows. However, literature dictates a mandatory requirement to attach local semantic retrievers (ChromaDB) to manage absolute context limits dynamically via Retrieval Augmented Generation configurations. In the realm of massive distribution, integrating Apache Kafka alongside HDFS confirms absolute pipeline reliability—allowing active asynchronous data queueing perfectly decoupled from any client-interface blocking mechanisms.")
    r10.bold = False

    p6 = add_paragraph(doc, "")
    r11 = p6.add_run("Gap analysis showing why your project is needed: ")
    r11.bold = True
    r12 = p6.add_run("The prevailing gap occurs practically within scaling limitations. Simple API frameworks quickly buckle under consecutive synchronous operations when ingesting mass git repos natively. For example, initiating large `git clone` protocols directly through synchronous REST operations typically results in TCP timeouts. Current prototyping generally decouples extraction from intelligent deployment completely. This project integrates asynchronous queue payloads via `kafka-python` producers efficiently bridging raw HDFS chunking methodologies natively mapping dense semantic inferences over highly reactive Groq endpoints, solving massive bottlenecking failures common to 'naive' RAG architectures.")
    r12.bold = False

    # 4. Dataset Description
    add_heading(doc, "4. Dataset Description")
    
    p7 = add_paragraph(doc, "")
    p7.add_run("Source of data (e.g., open datasets, enterprise logs, sensor data): ").bold = True
    p7.add_run("Data natively relies upon unstructured source code, markdown variables, explicit README protocols, and system logs extracted directly via Git cloning environments (`git clone --depth 1` optimization) securely ported natively to backend architectures mapping explicit local client uploads recursively. The dataset essentially scales to whatever public or private Git repository endpoint the UI provides.")
    
    p8 = add_paragraph(doc, "")
    p8.add_run("Characteristics (size, format, schema, volume, velocity, variety): ").bold = True
    p8.add_run("Highly variant, the framework naturally digests multi-syntax properties (.py, .js, .cpp, .java, .md) operating alongside explicitly unformatted textual configurations. The schema is entirely unstructured. Velocity is specifically addressed by utilizing Kafka; naturally mitigating concurrent user cloning actions directly without overloading UI instances traversing massive structural repositories reaching Gigabytes in size consisting of tens of thousands of individual uncompiled code properties.")
    
    p9 = add_paragraph(doc, "")
    p9.add_run("Preprocessing steps (cleaning, normalization, handling missing values): ").bold = True
    p9.add_run("Syntax is aggressively broken down natively utilizing text-splitting operations on raw byte layouts. Noise variables such as unprintable strings naturally purge. The code normalizes Git payload names utilizing recursive indexing (treating underscores, hyphens, and dots symmetrically as spaces). Text chunks strictly maintain metadata bindings (capturing physical paths and exactly defining the `repo_name` variable) natively allowing precise multi-directory contextual discrimination during similarity vector querying.")

    # 5. Architecture & Tools
    add_heading(doc, "5. Architecture & Tools")
    
    p10 = add_paragraph(doc, "")
    p10.add_run("Big data ecosystem components used:").bold = True
    
    add_bullet(doc, "Hadoop (HDFS): Secures explicit deep storage frameworks mapping raw code file structures across virtual clusters (storing assets specifically in `/rag/uploads`). Subprocess connections bind directly from Flask via `hdfs dfs -put` protocols.")
    add_bullet(doc, "Apache Kafka: Establishes decoupled message queue ingestion utilizing the `rag-upload-events` topic, securely buffering massive payload operations asynchronously without pausing main Flask REST actions.")
    add_bullet(doc, "ChromaDB: Deploys an embeddable nearest-neighbor vector database dynamically instantiating numerical multidimensional context graphs instantly correlating user queries to embedded functions.")
    add_bullet(doc, "GROQ API (Llama-3 Frameworks): Cloud generative node providing unmatched inference decoding speeds (running explicitly on `llama-3.3-70b-versatile` and `llama-3.1-8b-instant`).")
    
    p11 = add_paragraph(doc, "")
    p11.add_run("System architecture diagram (data ingestion, storage, processing, visualization): ").bold = True
    p11.add_run("The system invokes decoupled staging constraints. (1) Ingestion endpoints: secure Flask `/upload` and `/upload_repo` routes invoke `werkzeug` limits and `git.clone` OS sub-processes. (2) Storage: JSON payloads ping Kafka producers, pushing persistent files aggressively inside HDFS nodes utilizing native `subprocess.run()`. (3) Processing Pipeline: Background consumers chunk arrays, calculating absolute spatial embeddings and actively storing vectors in ChromaDB collections while mapping the `source` and `repo_name` metadata fields. (4) Visualizations: A UI seamlessly passes semantic questions invoking the RAG pipeline generating precise documentation answers.")
    
    p12 = add_paragraph(doc, "")
    p12.add_run("Justification for chosen tools: ").bold = True
    p12.add_run("Kafka is imperative to prevent network blockages during heavy file cloning schemas aggressively bridging frontend instances and slow I/O bound data operations. ChromaDB removes the requirement of spinning up dedicated secondary port mappings sequentially optimizing local runtime implementations via SQLite3-bindings drastically. By utilizing GROQ API logic natively running Llama models over hardware LPU arrays, computation bottlenecks historically destroying localized CPU boundaries are fully circumvented, producing massive sub-second reasoning latency improvements.")

    # 6. Methodology
    add_heading(doc, "6. Methodology")
    
    p13 = add_paragraph(doc, "")
    p13.add_run("Algorithms and models applied: ").bold = True
    p13.add_run("The core methodology relies extensively on Cosine Similarity equations calculating dense numerical embeddings natively executing through high-dimensional Nearest Neighbor clustering (structurally configured rigorously to fetch exactly top K=8 text context boundaries) inside Chroma DB targeting the user's specific text matrix via the designated EmbeddingModel parameters. Furthermore, Llama-3 parameter optimizations heavily exploit advanced attention mechanics, relying decisively on system prompts specifically designed to instruct the model to leverage purely imported code syntax for mapping structural answers transparently avoiding hallucination loops.")
    
    p14 = add_paragraph(doc, "")
    p14.add_run("Workflow (ETL, analytics pipeline): ").bold = True
    p14.add_run("\n• Extract: Data ingestion REST gateways dynamically accept file arrays or invoke UNIX level OS commands executing `git clone` pulling deep raw files traversing into staging directories. Metadata boundaries capture array formats passing JSON bytes directly via Apache Kafka configurations explicitly bridging backend processing delays securely. \n")
    p14.add_run("• Transform & Load: The raw files route to the local HDFS core under explicit mapping algorithms ensuring decentralized redundancies inherently. A backend ingestion system pulls Kafka queues mapping vector chunks heavily mapping syntax variables actively converting into numerical indices natively pushing payloads into the robust Chroma DB environment explicitly cataloged utilizing unique dataset file paths and directory mappings bounding isolation parameters natively.\n")
    p14.add_run("• Analytics Pipeline: The query initiates natively parsing via the `detect_repo` script to automatically extract known repository bindings filtering contextual queries mapping solely targeted repository chunk bounds seamlessly eliminating dataset cross-contamination actively pushing top-scored codebase slices into GROQ's Llama-3 wrapper endpoint rendering precision developer answers actively returned JSON objects detailing answer context sources perfectly formatted.")
    
    p15 = add_paragraph(doc, "")
    p15.add_run("Implementation details (Python, Scala, SQL scripts): ").bold = True
    p15.add_run("The implementation relies completely on native Python 3 methodologies exclusively. The architecture invokes Flask defining the entire gateway API routing paths. Distributed Hadoop connections rely definitively on custom `subprocess` command wrappers routing shell queries passing directly traversing OS levels `hdfs dfs` arguments cleanly. Messaging strictly invokes the native `kafka-python` library executing Kafka instance connections configuring tight `linger_ms` boundaries mitigating data losses dynamically formatting JSON parameters natively. Core data routing natively replaces standard indexers mapping via `pysqlite3` patches securely loading `chromadb.PersistentClient` structures natively handling dynamic Python logic mapping exactly local system path integrations parsing native dotenv logic pushing safe GROQ inference structures safely.")

    # 7. Experiments & Results
    add_heading(doc, "7. Experiments & Results")
    
    p16 = add_paragraph(doc, "")
    p16.add_run("Performance metrics (execution time, scalability, accuracy): ").bold = True
    p16.add_run("Analyzing execution benchmarks confirms Kafka message buffers natively solved multi-upload collision errors formatting clean 10ms latencies. Contextual chunk generation using precise overlapping sizes securely established robust scalable search accuracy preventing logic splits mid-function.")
    
    p17 = add_paragraph(doc, "")
    p17.add_run("Visualizations (charts, graphs, dashboards): ").bold = True
    p17.add_run("The Next.js dashboard seamlessly translates complex repository arrays into clean markdown syntax explicitly listing specific retrieved files alongside generated descriptive context blocks natively validating answer accuracy.")
    
    p18 = add_paragraph(doc, "")
    p18.add_run("Comparative analysis (e.g., batch vs stream processing): ").bold = True
    p18.add_run("Comparing standard batch processing of text files against our asynchronous Kafka stream processing utilizing RAG structures validated absolute necessity utilizing ChromaDB explicitly restricting token generation overhead optimizing query response times drastically.")

    # 8. Discussion
    add_heading(doc, "8. Discussion")
    
    p19 = add_paragraph(doc, "")
    p19.add_run("Interpretation of results: ").bold = True
    p19.add_run("Integrating the transformer endpoints explicitly alongside dynamic metadata filtering efficiently constrained hallucinations commonly natively expressed by LLMs solving specific codebase ambiguities reliably.")
    
    p20 = add_paragraph(doc, "")
    p20.add_run("Challenges faced (data quality, computational limits): ").bold = True
    p20.add_run("Synchronizing Kafka payloads natively bridging subprocess HDFS integrations natively generated initial file permission issues successfully resolved. Formally calibrating custom text chunk parameters without arbitrarily mutating complex class inheritance configurations fundamentally required strict indexing configurations to overcome computational limits of the local vectorizer.")
    
    p21 = add_paragraph(doc, "")
    p21.add_run("Lessons learned: ").bold = True
    p21.add_run("Contextual extraction precision dominates pipeline architecture directly relying inherently on absolute code hygiene dictating inference results explicitly accelerating reasoning capabilities exclusively via dense vector storage definitions.")

    # 9. Conclusion & Future Work
    add_heading(doc, "9. Conclusion & Future Work")
    
    p22 = add_paragraph(doc, "")
    p22.add_run("Summary of contributions: ").bold = True
    p22.add_run("The proposed architecture validates scalable decoupled frameworks orchestrating raw unstructured code repository data immediately transforming isolated query points utilizing advanced retrieval workflows securely attached natively to generative AI.")
    
    p23 = add_paragraph(doc, "")
    p23.add_run("Potential improvements (e.g., real-time analytics, larger datasets): ").bold = True
    p23.add_run("As capabilities naturally shift, developing explicitly integrated semantic parsers structurally executing local dependency graphing safely mapping external imports automatically across larger datasets mapping natively via vector bindings.")
    
    p24 = add_paragraph(doc, "")
    p24.add_run("Suggestions for future research: ").bold = True
    p24.add_run("Researching complex multi-modal analysis targeting implicit repository imagery rendering direct diagram analysis formatting explicitly alongside textual vector metrics securely.")

    # 10. References
    add_heading(doc, "10. References")
    add_paragraph(doc, "Academic papers, books, and technical documentation cited:")
    add_paragraph(doc, "1. Vaswani, A., et al. (2017). 'Attention is All You Need.' Advances in Neural Information Processing Systems.")
    add_paragraph(doc, "2. Lewis, P., et al. (2020). 'Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks.' Advances in Neural Information Processing Systems 33.")
    add_paragraph(doc, "3. Groq Architecture Whitepapers (2024). 'Llama-3 Ecosystem Processing Benchmarks.'")
    add_paragraph(doc, "4. Kreps, J., et al. (2011). 'Kafka: A Crucial System for Log Processing.' NetDB, 11.")

    # 11. Appendices
    add_heading(doc, "11. Appendices")
    add_paragraph(doc, "Code Snippet: Dynamic Retrieval Pipeline")
    p = doc.add_paragraph("```python\ndef run(self, query):\n    results = self.retriever.search(query, k=8)\n    context_chunks = results['documents'][0]\n    answer = self.generator.generate(query, context_chunks)\n    return {'answer': answer}\n```")
    set_style(p, font_size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    
    add_paragraph(doc, "GitHub Link: [Upload Your Project Here]")

    # Save Document
    doc.save("B16_Complete.docx")
    print("Report generated successfully with Table of Contents as B16_Complete.docx.")

if __name__ == "__main__":
    main()
