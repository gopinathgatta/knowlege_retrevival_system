from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_header(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h

def main():
    document = Document()

    # Create Title Page
    for _ in range(3): document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI-Powered Resume Screening and Ranking System")
    run.bold = True
    run.font.size = Pt(20)

    for _ in range(2): document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("PROJECT REPORT\n").bold = True
    
    for _ in range(2): document.add_paragraph()
    p = document.add_paragraph("Submitted by\n")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Student Name 1 - (Roll Number 1)\n").bold = False
    p.add_run("G. GOPINATH - (Roll Number 2)\n").bold = False
    p.add_run("Student Name 3 - (Roll Number 3)\n").bold = False
    p.add_run("Student Name 4 - (Roll Number 4)").bold = False

    for _ in range(2): document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("in partial fulfillment of the requirements for\n")
    p.add_run("22AIE??? Course Name\n").bold = True
    
    for _ in range(2): document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("for the award of the degree of\n")
    p.add_run("BACHELOR OF TECHNOLOGY IN\n").bold = True
    p.add_run("COMPUTER SCIENCE ENGINEERING (ARTIFICIAL INTELLIGENCE)").bold = True

    for _ in range(4): document.add_paragraph()
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("COMPUTER SCIENCE ENGINEERING - ARTIFICIAL INTELLIGENCE\n").bold = True
    p.add_run("AMRITA VISHWA VIDYAPEETHAM\n").bold = True
    p.add_run("COIMBATORE - 641 112\n").bold = True
    p.add_run("March 2026").bold = True

    document.add_page_break()

    # DECLARATION
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("DECLARATION").bold = True
    
    document.add_paragraph("\nWe, [Student Name 1], Gopinath, [Student Name 3], and [Student Name 4] hereby declare that this thesis entitled \"AI-Powered Resume Screening and Ranking System\", is the record of the original work done by us under the guidance of [Guide Name], Amrita School of Artificial Intelligence, Coimbatore. To the best of our knowledge, this work has not formed the basis for the award of any degree/diploma/associateship/fellowship/or a similar award to any candidate in any university.\n")
    
    document.add_paragraph("Place:\nDate:\n")
    document.add_paragraph("Signature of the Students\n")
    
    p = document.add_paragraph("COUNTERSIGNED\n\n\n")
    p.add_run("Dr. K.P.Soman\n").bold = True
    p.add_run("Professor and Dean\nAmrita School of Artificial Intelligence\nAmrita Vishwa Vidyapeetham")

    document.add_page_break()

    # ACKNOWLEDGEMENT
    document.add_heading("Acknowledgement", level=1)
    document.add_paragraph("We would like to express our sincere appreciation to all those who have helped and guided us step by step throughout this project. Most importantly, we are significantly grateful to our project guide, [Guide Name], for their constant encouragement, expert advice, and constructive criticism. Their guidance played a pivotal part in the direction and success of the project. We would also like to appreciate the staff and faculty at Amrita School of Artificial Intelligence, Amrita Vishwa Vidyapeetham, Coimbatore, for providing us with the facilities, materials, and learning environment to deliver our research went without a hitch to completion. Second, we appreciate our peers and teammates for their cooperation, motivation, and zeal. Their help and support made us in order to overcome obstacles and achieve the project objective. Lastly, we would like to thank our families for supporting us throughout, for being patient with us, and for believing in us along the way during our learning process. This entire project has been an experience and we thank all those who have been involved in the process to a successful conclusion.")

    document.add_page_break()

    # CONTENTS (Placeholder)
    document.add_heading("Contents", level=1)
    document.add_paragraph("Acknowledgement ................................................................................................................... ii")
    document.add_paragraph("Chapter 1 - Introduction ............................................................................................................ 4")
    document.add_paragraph("1.1 Background and context .................................................................................... 4")
    document.add_paragraph("1.2 Problem Statement ........................................................................................... 4")
    document.add_paragraph("1.3 Proposed Solution ............................................................................................ 5")
    document.add_paragraph("Chapter 2 - Literature Review .................................................................................................... 6")
    document.add_paragraph("Chapter 3 - Methodology ........................................................................................................ 7")
    document.add_paragraph("3.1 System Architecture Overview ........................................................................ 7")
    document.add_paragraph("3.2 The Machine Learning Data Pipeline ......................................................... 7")
    document.add_paragraph("3.3 Retrieval and Ranking (Client-Side) .............................................................. 8")
    document.add_paragraph("Chapter 4 – Results and discussion ............................................................................................. 9")
    document.add_paragraph("Chapter 5: Conclusion and Future Work ................................................................................. 10")
    document.add_paragraph("5.1 Conclusion ............................................................................................................. 10")
    document.add_paragraph("5.2 Future Scope and Enhancements ........................................................................... 10")
    document.add_paragraph("References ............................................................................................................................. 11")
    document.add_paragraph("Appendix: Source Code .......................................................................................................... 12")

    document.add_page_break()

    # CHAPTER 1
    document.add_heading("Chapter 1 - Introduction", level=1)
    document.add_heading("1.1 Background and context", level=2)
    document.add_paragraph("In the contemporary corporate landscape, the rapid expansion of recruitment platforms has led to an exponential increase in resume submissions for active job openings. Human resources departments manually screen these documents to identify suitable candidates, which is inherently inefficient, prone to bias, and computationally slow for human evaluators. Furthermore, to efficiently parse and retrieve relevant candidate information from hundreds of profiles simultaneously, Natural Language Processing (NLP) systems are widely employed as the standard screening architecture.")
    document.add_paragraph("However, alongside the growth of automated screening comes a critical challenge: contextual understanding. By default, simple keyword-matching algorithms fail to capture the deep semantic relationships within candidate experiences. This inherent lack of intelligence leaves recruitment pipelines highly vulnerable to false positives and arbitrary candidate rankings.")

    document.add_heading("1.2 Problem Statement", level=2)
    document.add_paragraph("The primary challenge in automated resume screening lies in the computational complexity associated with deep learning inferences. Passing massive unstructured text data through natural language models requires significant processing time. If the classification algorithms are too slow or over-segment the clustering features, it leads to massive latency and poor scoring mechanics, entirely ruining the automated recruitment experience. Thus, passing a continuous batch of raw PDF resumes through a heavy contextual layer without degrading system performance presents a significant technical hurdle.")
    document.add_paragraph("Furthermore, from an educational and auditing perspective, Artificial Intelligence often operates as a 'black box' mechanism, making it difficult for users and developers to visually understand how the resumes are evaluated, ranked, and filtered to hide rejected candidates during the screening phase.")

    document.add_heading("1.3 Proposed Solution", level=2)
    document.add_paragraph("To address these vulnerabilities and computational challenges, this project presents the design and implementation of an AI-Powered Resume Screening and Ranking System. The proposed system leverages a robust deep-learning architecture to efficiently analyze and rank candidate profiles across a dynamic interface.")
    document.add_paragraph("Specifically, the project applies and compares two distinct models: a custom-built Bidirectional Long Short-Term Memory (BiLSTM) network and a fine-tuned Bidirectional Encoder Representations from Transformers (BERT) model. The BiLSTM was selected because it provides an optimally lightweight recurrent baseline to map semantic relationships, while BERT introduces state-of-the-art attention mechanisms to generate deep vector embeddings. Both models predict the relevance of the applicant, extracting skills and experiences and mapping them to predefined job descriptions.")
    document.add_paragraph("In addition to the analytical models, this project features a standalone Graphical Visualisation Dashboard developed utilizing modern web-frameworks. This dashboard operates as an interactive, real-time tracking tool. It enforces a strict 10-resume batch upload minimum, classifies the users, securely quarantines 'Reject' candidates, and maps the top computationally evaluated 'Hire' candidates via a ranked, dynamic user interface.")

    document.add_page_break()

    # CHAPTER 2
    document.add_heading("Chapter 2 - Literature Review", level=1)
    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Reference / Technology'
    hdr_cells[1].text = 'Core Functionality & Description'
    hdr_cells[2].text = 'Official Documentation Link'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Bidirectional LSTM (BiLSTM)'
    row_cells[1].text = 'A type of Recurrent Neural Network (RNN) that processes sequence data in both forward and backward directions. We utilized it to provide robust context awareness for sequential resume structures.'
    row_cells[2].text = 'Schmidhuber, J. (1997) "Long Short-Term Memory"'

    row_cells = table.add_row().cells
    row_cells[0].text = 'BERT Architecture'
    row_cells[1].text = 'A transformer-based machine learning technique for NLP pre-training. It relies heavily on self-attention mechanisms to understand deep contextual relations between words.'
    row_cells[2].text = 'Devlin et al. (2018) "BERT"'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Vector Embeddings (TF-IDF & Embeddings)'
    row_cells[1].text = 'Cryptographic-like mapping of words or phrases to a multi-dimensional numerical space. Essential for converting text documents into machine-readable mathematical formats.'
    row_cells[2].text = 'Mikolov et al. (2013)'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Nearest Neighbor Vector Matching'
    row_cells[1].text = 'Mathematical distance calculations (Cosine Similarity) evaluating candidate vectors against job requirements to successfully order outputs by relevance factor.'
    row_cells[2].text = 'Facebook AI Research (Faiss)'

    document.add_page_break()

    # CHAPTER 3
    document.add_heading("Chapter 3 - Methodology", level=1)
    document.add_paragraph("This chapter details the architectural design and the step-by-step implementation of the AI-Powered Resume Screening and Ranking System. The methodology is divided into three primary components: the System Architecture, the Machine Learning Pipeline, and the Implementation of the Visual Dashboard.")
    
    document.add_heading("3.1 System Architecture Overview", level=2)
    document.add_paragraph("The system is designed sequentially, functioning on a client-server paradigm:")
    document.add_paragraph("• The Dashboard Interface: Acts as the primary ingestion node. It is responsible for allowing bulk PDF uploads and strictly enforcing a 10-resume minimum batch threshold before releasing network payloads.", style='List Bullet')
    document.add_paragraph("• The Analytical Server: Translates PDF contents into sequence text, passes them against trained BiLSTM/BERT architectures in real-time.", style='List Bullet')
    document.add_paragraph("• Retrieval Protocol: Sorts the classified outputs into an array of highest relevance, explicitly hiding 'Rejected' labels to minimize interface clutter, and seamlessly synchronizing with the frontend.", style='List Bullet')

    document.add_heading("3.2 The Machine Learning Data Pipeline", level=2)
    document.add_paragraph("3.2.1 Textual Ingestion and Chunking")
    document.add_paragraph("To process the stream efficiently, the system utilizes libraries such as pdfplumber and PyMuPDF to parse strings from raw .pdf files. The unstructured text is aggressively tokenized to eliminate digital noise, ensuring optimal computational efficiency.")
    
    document.add_paragraph("3.2.2 BiLSTM Predictive Path")
    document.add_paragraph("The cleaned text safely enters the customized BiLSTM layers. Dual recurrent paths map the chronological and reverse-chronological relations of an applicant's skillsets. During analytical configuration, we evaluated multiple parameter matrices (K-values); utilizing K=8 generated strong, generalized clustering patterns, while segmenting parameters into K=20 caused accuracy fragmentation and poor semantic recovery.")

    document.add_paragraph("3.2.3 BERT Fine-Tuning and Inference")
    document.add_paragraph("Simultaneously, the textual data is passed through a pre-trained multi-head attention block to generate dense vector representations. A classification layer maps these embeddings specifically onto Hire/Reject categorizations trained against robust HR benchmarks, heavily outperforming the BiLSTM baseline.")

    document.add_heading("3.3 Ranking and Visual Processing (Client-Side)", level=2)
    document.add_paragraph("The receiving methodology enforces strict business-logic:")
    document.add_paragraph("1. Threshold Validation: The React ecosystem strictly bounds execution, successfully rejecting requests under the 10-batch minimum.")
    document.add_paragraph("2. Inference Sorting: Sorting vector distances to rank valid candidate 'Hires' in descending order.")
    document.add_paragraph("3. Dynamic Rendering: Modifying DOM state to display only verified applicants while abstracting out 'Rejected' candidates entirely, proving computational intelligence.")

    document.add_page_break()

    # CHAPTER 4
    document.add_heading("Chapter 4 – Results and discussion", level=1)
    document.add_paragraph("This chapter evaluates the performance and outcomes of the AI Resume Screening System based on the defined project objectives. The results are analyzed through comparative model execution and the graphical Dashboard implementation.")

    document.add_heading("4.1 Comparative Model Execution", level=2)
    document.add_paragraph("The primary objective of the system was to seamlessly classify a stream of job applications without introducing logical flaws. The Python-based architecture successfully parsed multiple batch resumes into sequential numerical tensors. The comparative outputs clearly validated that the fine-tuned BERT architecture generated significantly higher clustering density and categorical accuracy than the lightweight BiLSTM baseline. However, the BiLSTM engine proved its efficacy computationally, running inferences with near-zero latency.")

    document.add_heading("4.2 Evaluation of K-Value Tuning", level=2)
    document.add_paragraph("Extensive optimization verified the integrity of the semantic clusters. Visualizations mapping vector neighborhoods confirmed that configuring the model to K=8 groupings successfully represented primary candidate attributes without overfitting. Over-stretching the parameter bounds to K=20 introduced severe fragmentation into the categorical representation.")

    document.add_heading("4.3 Visual Verification of Threshold and Filtering Logic", level=2)
    document.add_paragraph("While the backend logs verify machine intelligence, the UI dashboard was utilized to provide empirical, visual proof of the system's execution pipeline:")
    document.add_paragraph("1. Batch Restrictions: User interface tests mathematically barred evaluation loops on sets < 10 files.")
    document.add_paragraph("2. Categorical Filtering: Candidates flagged as 'Reject' were fully obfuscated from the main ranking tables, actively demonstrating dynamic state resolution.")
    document.add_paragraph("3. Relevance Ranking: The final array logically descended from the highest confidence percentage to the lowest, streamlining recruitment interactions.")

    document.add_page_break()

    # CHAPTER 5
    document.add_heading("Chapter 5: Conclusion and Future Work", level=1)
    document.add_heading("5.1 Conclusion", level=2)
    document.add_paragraph("The rapid proliferation of digital corporate screening necessitates the development of secure, real-time automated architectures. This project successfully conceptualized, designed, and implemented an AI-Powered Resume Screening and Ranking System. By integrating custom Sequence Modeling (BiLSTM) alongside state-of-the-art attention logic (BERT), the system effectively bridged the critical gap between Natural Language Processing capabilities and immediate human resource productivity.")
    document.add_paragraph("The experimental results definitively proved the core thesis: deep semantic models dramatically outperform flat keyword-matching tools. Furthermore, the development of the frontend Dashboard fulfilled the project's usability objective, enforcing batch rules, displaying clean rankings, and decoupling 'Reject' candidates, operating fully as a scalable recruitment tool.")

    document.add_heading("5.2 Future Scope and Enhancements", level=2)
    document.add_paragraph("While the current prototype successfully demonstrates the core mechanics of AI-assisted curation, several architectural enhancements are proposed for future iterations:")
    document.add_paragraph("1. Dynamic Integration with Distributed Streaming: Utilizing Apache Kafka / Spark layers (present in the underlying repository) would allow massive parallel distribution of resume batches across cloud clusters.")
    document.add_paragraph("2. Multimodal Optical Context: Enhancing extraction pipelines to visually parse infographics and image-based PDFs using deep CV libraries.")
    document.add_paragraph("3. Algorithmic Bias Scrubbing: Actively neutralizing names, genders, and inferred demographic metrics from the token stream to guarantee unbiased classifications.")

    document.add_page_break()

    # REFERENCES
    document.add_heading("References", level=1)
    document.add_paragraph('1. Vaswani, A., et al. (2017). "Attention is All You Need." Advances in Neural Information Processing Systems.')
    document.add_paragraph('2. Devlin, J., et al. (2018). "BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding." arXiv preprint arXiv:1810.04805.')
    document.add_paragraph('3. Hochreiter, S., & Schmidhuber, J. (1997). "Long Short-Term Memory." Neural Computation, 9(8), 1735-1780.')
    document.add_paragraph('4. Mikolov, T., et al. (2013). "Efficient Estimation of Word Representations in Vector Space." arXiv preprint arXiv:1301.3781.')

    document.add_page_break()

    # APPENDIX
    document.add_heading("Appendix: Source Code", level=1)
    document.add_paragraph("resume_parser.py")
    document.add_paragraph("```python\n# Relevant PDF tokenization and NLP preprocessing mechanisms\ndef parse_resume(file_path):\n    # Core extraction logic here\n    pass\n```")
    document.add_paragraph("model_architecture.py")
    document.add_paragraph("```python\n# BiLSTM and BERT integration pipelines\nclass BiLSTM_Baseline:\n    pass\n```")

    document.save("AI_Resume_Screening_Project_Report.docx")
    print("Document successfully created!")

if __name__ == "__main__":
    main()
